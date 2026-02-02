"""
Document Processor Service - Handle all document types
"""
import io
import logging
import magic
from typing import Tuple, List, Optional
from PIL import Image
from docx import Document as DocxDocument
from openpyxl import load_workbook

from app.services.pdf_service import pdf_service

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing various document types"""

    @staticmethod
    async def detect_file_type(file_blob: bytes, filename: str) -> str:
        """
        Detect file type using magic numbers and filename.

        Args:
            file_blob: File content as bytes
            filename: Original filename

        Returns:
            File type string (pdf, docx, xlsx, image, text, etc.)
        """
        try:
            mime = magic.from_buffer(file_blob, mime=True)

            # Map MIME types to our types
            if mime == "application/pdf":
                return "pdf"
            elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return "docx"
            elif mime in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"]:
                return "xlsx"
            elif mime.startswith("image/"):
                return "image"
            elif mime.startswith("text/"):
                return "text"
            else:
                # Fallback to extension
                if filename.lower().endswith(".pdf"):
                    return "pdf"
                elif filename.lower().endswith((".docx", ".doc")):
                    return "docx"
                elif filename.lower().endswith((".xlsx", ".xls")):
                    return "xlsx"
                elif filename.lower().endswith((".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff")):
                    return "image"
                elif filename.lower().endswith(".txt"):
                    return "text"
                else:
                    return "unknown"

        except Exception as e:
            logger.error(f"File type detection error: {e}")
            return "unknown"

    @staticmethod
    async def extract_content(file_blob: bytes, file_type: str) -> Tuple[Optional[str], List[bytes]]:
        """
        Extract text and images from document based on type.

        Args:
            file_blob: File content as bytes
            file_type: Detected file type

        Returns:
            Tuple of (text_content, list_of_image_bytes)
        """
        try:
            if file_type == "pdf":
                return await DocumentProcessor._extract_pdf(file_blob)
            elif file_type == "docx":
                return await DocumentProcessor._extract_docx(file_blob)
            elif file_type == "xlsx":
                return await DocumentProcessor._extract_xlsx(file_blob)
            elif file_type == "image":
                return await DocumentProcessor._extract_image(file_blob)
            elif file_type == "text":
                return await DocumentProcessor._extract_text(file_blob)
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return None, []

        except Exception as e:
            logger.error(f"Content extraction error for {file_type}: {e}")
            raise

    @staticmethod
    async def _extract_pdf(file_blob: bytes) -> Tuple[str, List[bytes]]:
        """Extract from PDF using PDFService"""
        return await pdf_service.extract_text_and_images(file_blob)

    @staticmethod
    async def _extract_docx(file_blob: bytes) -> Tuple[str, List[bytes]]:
        """Extract text and images from DOCX"""
        try:
            doc = DocxDocument(io.BytesIO(file_blob))

            # Extract text
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    table_text.append(row_text)
                text_parts.append("\n".join(table_text))

            full_text = "\n\n".join(text_parts)

            # Extract images (embedded in document)
            images = []
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_blob = rel.target_part.blob
                        images.append(image_blob)
                    except Exception as e:
                        logger.warning(f"Failed to extract DOCX image: {e}")

            logger.info(f"Extracted {len(full_text)} chars and {len(images)} images from DOCX")
            return full_text, images

        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise

    @staticmethod
    async def _extract_xlsx(file_blob: bytes) -> Tuple[str, List[bytes]]:
        """Extract text from Excel spreadsheet"""
        try:
            workbook = load_workbook(io.BytesIO(file_blob), data_only=True)

            text_parts = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"--- Sheet: {sheet_name} ---")

                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip():
                        text_parts.append(row_text)

            full_text = "\n".join(text_parts)

            logger.info(f"Extracted {len(full_text)} chars from XLSX with {len(workbook.sheetnames)} sheets")
            return full_text, []

        except Exception as e:
            logger.error(f"XLSX extraction error: {e}")
            raise

    @staticmethod
    async def _extract_image(file_blob: bytes) -> Tuple[str, List[bytes]]:
        """Extract image - return as image blob with minimal text"""
        try:
            # Verify it's a valid image
            img = Image.open(io.BytesIO(file_blob))
            width, height = img.size

            # Simple text description
            text = f"Image: {width}x{height} pixels, format: {img.format}"

            logger.info(f"Processed image: {text}")
            return text, [file_blob]

        except Exception as e:
            logger.error(f"Image extraction error: {e}")
            raise

    @staticmethod
    async def _extract_text(file_blob: bytes) -> Tuple[str, List[bytes]]:
        """Extract plain text"""
        try:
            # Try UTF-8 first, fallback to latin-1
            try:
                text = file_blob.decode('utf-8')
            except UnicodeDecodeError:
                text = file_blob.decode('latin-1')

            logger.info(f"Extracted {len(text)} chars from text file")
            return text, []

        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            raise

    @staticmethod
    async def optimize_image_for_vision_api(image_blob: bytes, max_size_kb: int = 500) -> bytes:
        """
        Optimize image for Vision API by resizing if too large.

        Args:
            image_blob: Original image bytes
            max_size_kb: Maximum size in KB

        Returns:
            Optimized image bytes
        """
        try:
            current_size_kb = len(image_blob) / 1024

            if current_size_kb <= max_size_kb:
                return image_blob

            # Resize image
            img = Image.open(io.BytesIO(image_blob))

            # Calculate new dimensions
            scale_factor = (max_size_kb / current_size_kb) ** 0.5
            new_width = int(img.width * scale_factor)
            new_height = int(img.height * scale_factor)

            # Resize
            img_resized = img.resize((new_width, new_height), Image.Resampling.LANCZOS)

            # Save to bytes
            output = io.BytesIO()
            img_resized.save(output, format=img.format or "JPEG", quality=85)
            optimized_blob = output.getvalue()

            logger.info(f"Optimized image from {current_size_kb:.1f}KB to {len(optimized_blob)/1024:.1f}KB")
            return optimized_blob

        except Exception as e:
            logger.warning(f"Image optimization failed: {e}, using original")
            return image_blob


# Global instance
document_processor = DocumentProcessor()
